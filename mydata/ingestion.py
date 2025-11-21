"""Unified ingestion pipeline for all data sources"""

import hashlib
from pathlib import Path
from typing import Optional, List
from uuid import UUID, uuid4
from datetime import datetime
from sqlmodel import Session
from .models import Document, Chunk
from .storage import EncryptedStorage
from .embedder import Embedder
from .vectordb import VectorDB
from .ml_organizer import MLOrganizer


class IngestionPipeline:
    """Handles ingestion from all sources: files, emails, stdin"""

    def __init__(
        self,
        db_session: Session,
        storage: EncryptedStorage,
        embedder: Embedder,
        vectordb: VectorDB,
        ml_organizer: Optional[MLOrganizer] = None,
    ):
        self.db = db_session
        self.storage = storage
        self.embedder = embedder
        self.vectordb = vectordb
        self.ml_organizer = ml_organizer

    def ingest_file(self, file_path: Path) -> Optional[UUID]:
        """Ingest a file (supports TXT, PDF, DOCX, XLSX, CSV, JSON, MD)"""
        if not file_path.exists():
            print(f"✗ File not found: {file_path}")
            return None

        # Read file
        try:
            content = file_path.read_bytes()
        except Exception as e:
            print(f"✗ Error reading file: {e}")
            return None

        # Extract text based on file type
        text = self._extract_text_from_file(file_path, content)
        if text is None:
            print(f"⚠ Skipping unsupported file: {file_path.name}")
            return None

        if not text.strip():
            print(f"⚠ Skipping empty file: {file_path.name}")
            return None

        # Compute hash for deduplication
        file_hash = self.storage.compute_hash(content)

        # Check if already exists
        existing = self.db.query(Document).filter(Document.file_hash == file_hash).first()
        if existing:
            print(f"⚠ File already indexed: {file_path.name}")
            return existing.id

        # Extract file metadata for change tracking
        from .file_metadata import get_file_metadata
        metadata = get_file_metadata(file_path)

        # Create document
        doc = Document(
            source=f"file://{file_path}",
            source_type="file",
            mime_type=self._detect_mime_type(file_path),
            file_hash=file_hash,
            raw_text=text,
            file_modified_at=metadata.get('modified_at'),
            file_created_at=metadata.get('created_at'),
            file_size_bytes=metadata.get('size_bytes'),
            file_owner=metadata.get('owner'),
        )

        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        # Process chunks and embeddings
        self._process_document(doc, text)

        print(f"✓ Ingested: {file_path.name} (ID: {str(doc.id)[:8]}...)")
        return doc.id

    def ingest_text(self, text: str, source: str = "stdin") -> Optional[UUID]:
        """Ingest raw text (from paste or other sources)"""
        # Check for semantic duplicates
        if self._is_semantic_duplicate(text):
            print(f"⚠ Semantically similar content already exists (skipping)")
            return None

        # Create document
        doc = Document(
            source=source,
            source_type="paste" if source == "stdin" else "other",
            raw_text=text,
        )

        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        # Process chunks and embeddings
        self._process_document(doc, text)

        print(f"✓ Ingested text (ID: {str(doc.id)[:8]}...)")
        return doc.id

    def ingest_email(self, email_data: dict) -> Optional[UUID]:
        """Ingest an email"""
        from datetime import datetime

        timestamp = datetime.now().strftime("%H:%M:%S")
        text = email_data["full_text"]
        source = f"email://{email_data.get('sender', 'unknown')}/{email_data.get('uid', 'unknown')}"

        # Check for semantic duplicates (emails can be forwarded/duplicated)
        if self._is_semantic_duplicate(text, threshold=0.98):
            subject = email_data.get("subject", "(no subject)")[:40]
            print(f"⚠ [{timestamp}] Duplicate email skipped: {subject}...")
            return None

        # Create document
        doc = Document(
            source=source,
            source_type="email",
            raw_text=text,
        )

        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        # Process chunks and embeddings
        chunk_count = self._process_document(doc, text)

        subject = email_data.get("subject", "(no subject)")[:50]
        sender = email_data.get("sender", "unknown")[:30]
        print(f"✓ [{timestamp}] Email ingested: '{subject}'")
        print(f"   • From: {sender}")
        print(f"   • Document ID: {str(doc.id)[:8]}...")
        print(f"   • Chunks created: {chunk_count}")
        print(f"   • Size: {len(text)} chars")
        return doc.id

    def _process_document(self, doc: Document, text: str) -> int:
        """Process document: chunk, embed, and organize"""
        # Simple chunking (split by paragraphs or fixed size)
        chunks = self._chunk_text(text, max_length=512)

        # Create chunk records and embeddings
        chunk_ids = []
        chunk_texts = []

        for i, chunk_text in enumerate(chunks):
            chunk = Chunk(
                doc_id=doc.id,
                text=chunk_text,
                start_offset=i * 512,  # Approximate
                end_offset=i * 512 + len(chunk_text),
            )
            self.db.add(chunk)
            chunk_ids.append(chunk.id)
            chunk_texts.append(chunk_text)

        self.db.commit()

        # Generate embeddings
        if chunk_texts:
            embeddings = self.embedder.embed_batch(chunk_texts)

            # Store in vector DB
            for chunk_id, embedding in zip(chunk_ids, embeddings):
                self.vectordb.upsert(
                    doc_id=str(chunk_id),
                    vector=embedding,
                    payload={
                        "doc_id": str(doc.id),
                        "text": chunk_texts[chunk_ids.index(chunk_id)][:200],
                        "source": doc.source,
                    },
                )

        # Run ML organization
        if self.ml_organizer:
            self.ml_organizer.organize_document(str(doc.id), text)

        return len(chunks)

    def _chunk_text(self, text: str, max_length: int = 512) -> List[str]:
        """Simple text chunking"""
        # Split by paragraphs first
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) < max_length:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks if chunks else [text[:max_length]]

    def _is_semantic_duplicate(self, text: str, threshold: float = 0.95) -> bool:
        """Check if semantically similar document already exists"""
        # Use first 500 chars as sample for speed
        sample = text[:500] if len(text) > 500 else text

        try:
            # Embed the sample
            query_emb = self.embedder.embed(sample)

            # Search for similar documents
            results = self.vectordb.search(
                query_vector=query_emb,
                limit=1,
                score_threshold=threshold
            )

            return len(results) > 0

        except Exception as e:
            # If check fails, allow ingestion (fail open)
            print(f"⚠ Semantic dedup check failed: {e}")
            return False

    def _extract_text_from_file(self, file_path: Path, content: bytes) -> Optional[str]:
        """
        Extract text from various file types.

        🔒 READ-ONLY OPERATION - This method ONLY reads files, never modifies them.

        Supported formats:
        - Text: .txt, .md, .csv, .json, .log
        - PDF: .pdf (using pypdf)
        - Word: .docx (using python-docx)
        - Excel: .xlsx, .xls (using openpyxl)

        Returns:
            Extracted text or None if unsupported/error
        """
        suffix = file_path.suffix.lower()

        try:
            # Plain text files
            if suffix in ['.txt', '.md', '.csv', '.json', '.log']:
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try with different encoding
                    try:
                        return content.decode('latin-1')
                    except Exception:
                        return None

            # PDF files
            elif suffix == '.pdf':
                try:
                    from pypdf import PdfReader
                    from io import BytesIO

                    pdf_file = BytesIO(content)
                    reader = PdfReader(pdf_file)

                    text_parts = []
                    for page_num, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        except Exception as e:
                            print(f"⚠ Error extracting page {page_num + 1}: {e}")
                            continue

                    if text_parts:
                        return '\n\n'.join(text_parts)
                    else:
                        print(f"⚠ No text extracted from PDF (may be scanned/image-based)")
                        return None

                except Exception as e:
                    print(f"⚠ PDF extraction failed: {e}")
                    return None

            # Word documents (.docx)
            elif suffix == '.docx':
                try:
                    from docx import Document as DocxDocument
                    from io import BytesIO

                    docx_file = BytesIO(content)
                    doc = DocxDocument(docx_file)

                    # Extract paragraphs
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

                    # Extract tables
                    table_texts = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = '\t'.join([cell.text for cell in row.cells])
                            if row_text.strip():
                                table_texts.append(row_text)

                    all_text = '\n\n'.join(paragraphs)
                    if table_texts:
                        all_text += '\n\n' + '\n'.join(table_texts)

                    return all_text if all_text.strip() else None

                except Exception as e:
                    print(f"⚠ DOCX extraction failed: {e}")
                    return None

            # Excel files (.xlsx, .xls)
            elif suffix in ['.xlsx', '.xls']:
                try:
                    from openpyxl import load_workbook
                    from io import BytesIO

                    excel_file = BytesIO(content)
                    wb = load_workbook(excel_file, data_only=True, read_only=True)

                    text_parts = []
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        text_parts.append(f"=== Sheet: {sheet_name} ===")

                        for row in sheet.iter_rows(values_only=True):
                            row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                            if row_text.strip():
                                text_parts.append(row_text)

                    wb.close()
                    return '\n'.join(text_parts) if text_parts else None

                except Exception as e:
                    print(f"⚠ Excel extraction failed: {e}")
                    return None

            # Unsupported file type
            else:
                return None

        except Exception as e:
            print(f"⚠ Unexpected error extracting {file_path.name}: {e}")
            return None

    def _detect_mime_type(self, file_path: Path) -> str:
        """Detect MIME type from file extension"""
        suffix = file_path.suffix.lower()
        mime_types = {
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".json": "application/json",
            ".csv": "text/csv",
            ".log": "text/plain",
        }
        return mime_types.get(suffix, "application/octet-stream")
